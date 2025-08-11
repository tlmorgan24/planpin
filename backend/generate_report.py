from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image # installed as pip install Pillow
import fitz # installed as pip install pymupdf
from io import BytesIO
from get_cloud_data import get_data

def generate_report(access_token, refresh_token, user_id, plan_id, priority_limit=5, include_caption=False):

    MARKER_PATH = "marker.png"

    # max dimensions of images (inches):
    MAX_WIDTH = 5
    MAX_HEIGHT = 5

    marker_records, priority_marker_records, marker_images, plan_pdf_stream = get_data(access_token, refresh_token, user_id, plan_id, priority_limit)

    # Start from blank template with styles & footer defined:
    doc = Document("template.docx")
    # Template colours:
    # Primary colour: black
    # Secondary colour: #248E6B

    ## ---- Title & subtitle ----

    doc.add_heading("PlanPin report", level=0)
    doc.add_paragraph("This inspection report was generated with PlanPin", style="Subtitle")


    ## ---- Executive summary ----

    doc.add_heading("Executive summary", level=1)

    # Highest-severity items:
    doc.add_heading("Highest-severity items", level=2)

    num_priority_records = len(priority_marker_records)
    table = doc.add_table(rows=num_priority_records+1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Bold header row:
    headers = ['Item reference', 'Category', 'Severity']
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True

    for i, record in enumerate(priority_marker_records, 1): # start enumerating at 1 instead of 0
        table.cell(i, 0).text = str(record['reference']) or '! Not provided'
        table.cell(i, 1).text = record['category_name'] or '! Not provided'
        table.cell(i, 2).text = str(record['severity']) if record['severity'] is not None else '! Not provided' # need this logic, as otherwise would return the string 'None' if None type


    # Number of defects for each category:
    doc.add_heading("Summary statistics", level=2)

    unique_categories = {record['category_name'] or '! Not provided' for record in marker_records} # set comprehension automatically ignores duplicates
    # ^ Note category_name may be null, hence these are replaced with '! Not provided' (exclamation mark meaning will appear at start of sorted list)
    categories_list = sorted(unique_categories) # list sorted alphabetically

    table = doc.add_table(rows=len(categories_list)+1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Bold header row:
    headers = ['Category', 'Number of items', 'Maximum severity']
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
    
    for i, category in enumerate(categories_list, 1): # start enumerating at 1 instead of 0
        if category == '! Not provided':
            category = None # convert category back to None (having been previously set to '! Not provided' for sorting purposes)
        table.cell(i, 0).text = category or '! Not provided'
        relevant_marker_records = [record for record in marker_records if record['category_name'] == category]
        table.cell(i, 1).text = str( len(relevant_marker_records) )
        # Note, we have to filter out null values and return a default value if iteratable is empty (otherwise, either case will cause error):
        table.cell(i, 2).text = str( max((record['severity'] for record in relevant_marker_records if record['severity'] is not None), default='! Not provided') )


    ## ---- Full item data ----

    doc.add_heading("Full item data", level=1)

    for record in marker_records:

        marker_id = record['id']
        images = marker_images[marker_id]

        pdf_start_row = 4 # 0-based. PDF image is shown after text rows
        if include_caption:
            image_start_row = pdf_start_row + 2 # 0-based. images start after text rows, then PDF row, then PDF caption
            num_rows = pdf_start_row + 2 + 2*len(images) # total (not just image) rows. +1 because pdf_start_row is base-0, +1 again because of PDF caption row 
        else:
            image_start_row = pdf_start_row + 1 # 0-based. images start after text rows, then PDF row
            num_rows = pdf_start_row + 1 + len(images) # total (not just image) rows. +1 because pdf_start_row is base-0
        
        table = doc.add_table(rows=num_rows, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Bold defect reference as header (merged cell):
        header_1 = table.cell(0, 0)
        header_2 = table.cell(0, 1)
        header = header_1.merge(header_2)
        p = header.paragraphs[0]
        run = p.add_run(f"Item reference: #{record['reference'] or 'N/A'}")
        run.bold = True

        table.cell(1,0).text = "Category"
        table.cell(1,1).text = record['category_name'] or '! Not provided' # if category is null, set to "! Not provided" (note, setting to None object would cause error)
        
        table.cell(2,0).text = "Description"
        table.cell(2,1).text = record['description'] or '! Not provided'

        table.cell(3,0).text = "Severity"
        table.cell(3,1).text = str(record['severity']) if record['severity'] is not None else '! Not provided' # need this logic, as otherwise would return the string 'None' if None type

        for i in range(1,4):
            left_cell = table.cell(i, 0)
            left_cell.width = Inches(1.2) # right cell automatically changes so page width is still filled

        # PDF

        page_number = record['page_number']
        x = record['x']
        y = record['y']
        color = record['color'] # note color is already associated with marker as part of get_data function, even though in my database, color is part of the categories table (not markers table)
        marked_image = mark_plan(MARKER_PATH, plan_pdf_stream, page_number, x, y, color)
        paragraph = create_merged_paragraph(table, pdf_start_row, align='center')
        insert_image(marked_image, paragraph, MAX_WIDTH, MAX_HEIGHT)
        if include_caption:
            paragraph = create_merged_paragraph(table, pdf_start_row+1, style='Caption')
            run = paragraph.add_run("Item location on plan")

        # IMAGES

        if len(images) == 0: continue # do not try to create image rows if there are no images; skip to the next record

        if include_caption:
            image_rows = range(image_start_row, image_start_row - 1 + 2*len(images), 2) # e.g. if 1 image, image_rows=[image_start_row]; if 2 images, image_rows=[image_start_row, image_start_row + 2]; etc.
        else:
            image_rows = range(image_start_row, image_start_row + len(images))

        for row_index, image in zip(image_rows, images):
            paragraph = create_merged_paragraph(table, row_index, align='center') # for merged cell with centred image
            with Image.open(image) as img: # get image as Pillow image
                insert_image(img, paragraph, MAX_WIDTH, MAX_HEIGHT)

            if include_caption:
                paragraph = create_merged_paragraph(table, row_index+1, style='Caption')
                run = paragraph.add_run("This is a caption")

        doc.add_paragraph("Will this work?", style="Normal") # blank line after each paragraph (uses non-breaking space to be sure won't be treated as empty and tables merged)

    return doc


def create_merged_paragraph(table, row_index, align=None, style=None):
    cell1 = table.cell(row_index, 0)
    cell2 = table.cell(row_index, 1)
    merged_cell = cell1.merge(cell2)
    paragraph = merged_cell.paragraphs[0]
    if align == 'centre' or align == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if style:
        paragraph.style = style
    return paragraph


# Image is Pillow image, max width and max height are in inches:
def insert_image(image, paragraph, max_width, max_height):
    
    width_px, height_px = image.size
    dpi = image.info.get('dpi', (72, 72))  # default DPI if not present
    width_in = width_px / dpi[0]
    height_in = height_px / dpi[1]

    # Scale down if larger than max dimensions:
    scale = min(max_width / width_in, max_height / height_in, 1.0)
    new_width = Inches(width_in * scale)
    new_height = Inches(height_in * scale)

    # Convert to bytes in memory (as can't pass Pillow object directly to python-docx add_picture method):
    image_stream = BytesIO()
    image.save(image_stream, format="PNG") # convert image to bytes in memory 
    image_stream.seek(0) # rewind for reading

    # Insert:
    run = paragraph.add_run()
    run.add_picture(image_stream, width=new_width, height=new_height)


# Returns marked plan as Pillow Image object (note python-docx can't insert PDF or SVG, so we have to convert to regular image):
def mark_plan(marker_path, pdf_stream, page_number, x, y, color):

    doc = fitz.open(stream=pdf_stream, filetype='pdf')
    page = doc.load_page(page_number - 1)
    
    # Get page dimensions:
    rect = page.rect
    width = rect.width
    height = rect.height
    scale = min(width, height) / 595.28 # 595.28 is width of A4 page in pt, noting that marker png is originally sized to look appropriate on an A4 page (so we will want to scale it up by this amount)
    
    # Convert PDF to Pillow Image (required to then allow overlay of pin marker):
    zoom = 1/scale # such that 1 PDF point will map to "zoom" IMAGE pixels (not device pixels or CSS pixels). We divide by scale so that a big PDF is given just as many pixels as a small PDF (as the image will ultimately be the same size anyway)
    matrix = fitz.Matrix(zoom, zoom)
    pixmap = page.get_pixmap(matrix=matrix) # get pixmap (image) from PDF
    pdf_image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples) # get Pillow Image object from pixmap
    
    # Get marker image of appropriate color and size (note, width and height here are for the marker's rectangular bounding box, so doesn't matter than the marker itself is not a rectangle):
    marker_image = Image.open(marker_path).convert("RGBA")
    marker_image = recolor_marker(marker_image, color)
    w, h = marker_image.size
    new_size = (int(w * scale), int(h * scale)) # if scale relative to A4 page is bigger, marker will be bigger, and vice versa
    marker_image = marker_image.resize(new_size, resample=Image.LANCZOS)

    # As we know the scale of the pixmap is created from the "zoom" variable, we can immediately convert x and y (pt) to image pixels (px):
    x_px = x * zoom
    y_px = y * zoom

    # If marker is too close to edge, may be clipped off. So, add padding around the PDF image equal to marker size to prevent this:
    padding_left = new_size[0] # we will also make padding_right equal to this
    padding_top = new_size[1] # we will also make padding_bottom equal to this
    padded_canvas = Image.new("RGB", (pdf_image.width + 2*padding_left, pdf_image.height + 2*padding_top), color=(255, 255, 255))
    padded_canvas.paste(pdf_image, (padding_left, padding_top))
    
    # Overlay marker image onto PDF:
    # Marker is pasted according to top-left corner of bounding box. Have to modify x and y, as these specify the BOTTOM MIDDLE of bounding box.
    # Also have to modify x and y further, to reflect the padding of the padded_canvas
    # Also note paste location must be tuple of INTEGERS.
    x_px_topleft = int(x_px - new_size[0] / 2 + padding_left)
    y_px_topleft = int(y_px - new_size[1] + padding_top)
    padded_canvas.paste(marker_image, (x_px_topleft, y_px_topleft), marker_image)

    return padded_canvas


# color may be string hex code or RGBA tuple:
def recolor_marker(marker_image, new_color):

    # If no color (e.g. user has not selected category for that marker), default to #9cc7b8 "--mid-accent-color" (hardcoded, could improve to use one source of truth in future):
    if not new_color:
        new_color = '#9cc7b8'

    # Convert hex to RGB if needed
    if isinstance(new_color, str):
        new_color = hex_to_rgb(new_color)

    datas = marker_image.getdata()

    new_data = []
    for item in datas:
        # Detect white (with threshold tweaked for anti-alising):
        if item[0] > 200 and item[1] > 200 and item[2] > 200 and item[3] > 0:
            new_data.append((*new_color, item[3])) # preserve alpha
        else:
            new_data.append(item)

    marker_image.putdata(new_data)
    return marker_image

def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))